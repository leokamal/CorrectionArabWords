import os
import tempfile
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from fastapi import HTTPException, BackgroundTasks, UploadFile
from correction_words_service import generate_query
from typing import List
import mammoth
import json


def save_temp_file(file: UploadFile) -> str:
    """Save uploaded file temporarily and return its path."""
    if file.filename.endswith((".txt", ".doc", ".docx")):
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[-1]) as temp_file:
            temp_file.write(file.file.read())
            return temp_file.name
    raise HTTPException(status_code=400, detail="Unsupported file format. Use .txt, .doc, or .docx")

def read_file_content(file_path: str) -> str:
    """Read content from TXT, DOC, or DOCX file."""
    if file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file_path.endswith(".doc"):
        with open(file_path, "rb") as f:
            result = mammoth.extract_raw_text(f)
        return result.value
    return ""


def add_paragraph_rtl(doc, text):
    """Add a right-to-left (RTL) paragraph to the document."""
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Align text to the right

    # Add RTL (bidi) property to the paragraph
    paragraph_properties = paragraph._p.get_or_add_pPr()  # Access paragraph properties
    bidi = OxmlElement('w:bidi')  # Create the bidirectional element
    bidi.set(qn('w:val'), '1')  # Enable RTL
    paragraph_properties.append(bidi)

def process_and_save(file_path: str, selected_options: List[int]):
    """Process text, apply corrections, and save output as DOCX."""
    text = read_file_content(file_path)
    response = generate_query(selected_options, text)
    # Step 1: Remove unwanted characters like "```json" and backticks
    cleaned_response = response.replace("```json", "").replace("```", "").strip()

    # Step 2: Decode the cleaned response string if it's byte-encoded (optional in Python)
    decoded_response = cleaned_response.encode('utf-8').decode('utf-8')

    # Step 3: Parse the JSON response
    json_response = json.loads(decoded_response)
    corrected_text = json_response.get("corrected_text", "")
    details = json_response.get("details", [])

    # Save output to DOCX
    output_path = file_path.replace(os.path.splitext(file_path)[-1], "_corrected.docx")
    doc = Document()
    doc.add_heading("النص الأصلي", level=1)
    add_paragraph_rtl(doc, text)
    doc.add_heading("النص المصحح", level=1)
    add_paragraph_rtl(doc, corrected_text)
    
    doc.add_heading("سجل التعديلات", level=1)
    for detail in details:
        add_paragraph_rtl(doc, detail)
    
    doc.save(output_path)
    return os.path.basename(output_path)
